# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(11)

# 제목
title = doc.add_heading('READING JESUS DAILY Q.T. 사용 가이드', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = subtitle.add_run('"매일, 삶 속에서 예수님을 읽어갑니다"')
run_sub.italic = True
run_sub.font.size = Pt(14)

doc.add_paragraph()

# 카드 1
doc.add_heading('카드 1. Q.T.의 목적', level=1)
p = doc.add_paragraph()
p.add_run("이 큐티책은 '많이 읽기'를 넘어,\n말씀 앞에서 ")
run_bold = p.add_run("예배자로 응답(Respond)")
run_bold.bold = True
p.add_run("하는 하루를 돕기 위해 준비되었습니다.")

doc.add_paragraph()
box = doc.add_paragraph()
run_theme = box.add_run("1월 주제: 예배")
run_theme.bold = True

doc.add_paragraph("예배는 주일 한 시간으로 끝나는 행위가 아니라,\n말씀을 듣고 마음과 삶으로 응답하는\n하나님이 주신 삶의 리듬입니다.")

doc.add_paragraph()
doc.add_paragraph('─' * 40)
doc.add_paragraph()

# 카드 2
doc.add_heading('카드 2. 7분 버전 (초간단)', level=1)
p2_sub = doc.add_paragraph()
run_7min = p2_sub.add_run('"7분으로 시작합니다"')
run_7min.italic = True

table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = table.rows[0].cells
headers[0].text = '단계'
headers[1].text = '할 일'

data = [
    ('1', '본문 읽고 마음에 남는 단어에 동그라미'),
    ('2', '오늘의 필사 구절 한 줄 필사'),
    ('3', '묵상 길잡이에서 인사이트 한 문장 밑줄'),
    ('4', '묵상 질문에 한 문장 답'),
    ('5', '짧은 기도 + 하루 점검 체크'),
]

for i, (step, task) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = step
    row[1].text = task

doc.add_paragraph()
doc.add_paragraph('─' * 40)
doc.add_paragraph()

# 카드 3
doc.add_heading('카드 3. 24분 버전 (깊어지는)', level=1)
p3_sub = doc.add_paragraph()
run_24min = p3_sub.add_run('"24분으로 깊어집니다"')
run_24min.italic = True

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = table2.rows[0].cells
headers2[0].text = '단계'
headers2[1].text = '이름'
headers2[2].text = '설명'

data2 = [
    ('1', '내 말로 정리', '본문 2번 읽고 한 문장으로 요약'),
    ('2', '새로운 관점', '길잡이 본 후 본문 다시 읽기'),
    ('3', '삶으로 연결', '질문 곱씹으며 구체적 삶에 적용'),
    ('4', '오늘 기록', '은혜를 구체적 사건으로 기록'),
    ('5', '기도 응답', '묵상을 기도로 전환 + 한 문장 기도'),
    ('6', '나의 고백', '공동체와 나눌 한 문장 (QR로 나눔)'),
    ('7', '하루 점검', '몰입도 스스로 체크'),
]

for i, (step, name, desc) in enumerate(data2, 1):
    row = table2.rows[i].cells
    row[0].text = step
    row[1].text = name
    row[2].text = desc

doc.add_paragraph()
doc.add_paragraph('─' * 40)
doc.add_paragraph()

# 카드 4
doc.add_heading('카드 4. 월간 도전', level=1)
p4 = doc.add_paragraph()
p4.add_run("한 달 마무리 시 Q.T. 기록 가져오면\n")
run_3d = p4.add_run("3D 프린팅 글자")
run_3d.bold = True
p4.add_run(" 기념품 증정!")

doc.add_paragraph()

table3 = doc.add_table(rows=2, cols=6)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

months1 = ['1월', '2월', '3월', '4월', '5월', '6월']
letters1 = ['R', 'E', 'A', 'D', 'I', 'N']

for i, (m, l) in enumerate(zip(months1, letters1)):
    table3.rows[0].cells[i].text = m
    table3.rows[1].cells[i].text = l

doc.add_paragraph()

table4 = doc.add_table(rows=2, cols=6)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

months2 = ['7월', '8월', '9월', '10월', '11월', '12월']
letters2 = ['G', 'J', 'E', 'S', 'U', 'S']

for i, (m, l) in enumerate(zip(months2, letters2)):
    table4.rows[0].cells[i].text = m
    table4.rows[1].cells[i].text = l

doc.add_paragraph()
p_complete = doc.add_paragraph()
p_complete.add_run("12개월 모으면 ")
run_rj = p_complete.add_run("READING JESUS")
run_rj.bold = True
p_complete.add_run(" 완성!")

doc.add_paragraph()
p_imp = doc.add_paragraph()
run_imp = p_imp.add_run("중요한 것은:")
run_imp.bold = True

doc.add_paragraph("• 완벽하게 다 채운 사람만의 보상이 아닙니다")
doc.add_paragraph("• 조금 빈 곳이 있어도 괜찮습니다")
doc.add_paragraph("• 중요한 것은 멈췄다가도 다시 시작하는 흔적입니다")

doc.add_paragraph()
doc.add_paragraph('─' * 40)
doc.add_paragraph()

# 카드 5
doc.add_heading('카드 5. 핵심 메시지', level=1)
quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_quote = quote.add_run('"목표는 완독이 아니라,\n매일 주님께 응답하는 예배의 지속입니다"')
run_quote.bold = True
run_quote.font.size = Pt(14)

doc.add_page_break()

# 앱 섹션
app_title = doc.add_heading('앱으로 묵상 나누기', 0)
app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 카드 6
doc.add_heading('카드 6. STEP 1 - QR 스캔', level=1)
p6 = doc.add_paragraph()
p6.add_run("큐티책의 ")
run_qr = p6.add_run("QR코드")
run_qr.bold = True
p6.add_run("를 스캔하세요")

doc.add_paragraph()
doc.add_paragraph("📷 → 리딩지저스 묵상 나눔 페이지로 이동")

doc.add_paragraph()
doc.add_paragraph('─' * 40)
doc.add_paragraph()

# 카드 7
doc.add_heading('카드 7. STEP 2 - 묵상 등록', level=1)
p7 = doc.add_paragraph()
run_confess = p7.add_run('"나의 고백"')
run_confess.bold = True
p7.add_run(" 한 문장을 입력하세요")

doc.add_paragraph()
example = doc.add_paragraph()
run_ex = example.add_run('예시: "주님, 오늘도 내 길을 인도하소서"')
run_ex.italic = True

doc.add_paragraph()
doc.add_paragraph("• 길게 쓰지 않아도 됩니다")
doc.add_paragraph("• 오늘 말씀을 통과한 한 문장이면 충분합니다")

doc.add_paragraph()
doc.add_paragraph('─' * 40)
doc.add_paragraph()

# 카드 8
doc.add_heading('카드 8. STEP 3 - 함께 나눔', level=1)
p8 = doc.add_paragraph()
p8.add_run("등록된 묵상은 ")
run_share = p8.add_run("교회 공동체와 공유")
run_share.bold = True
p8.add_run("됩니다")

doc.add_paragraph()
doc.add_paragraph("• 다른 성도님들의 고백도 읽어보세요")
p_final = doc.add_paragraph()
p_final.add_run("• ")
run_final = p_final.add_run("한 문장이 공동체를 살립니다")
run_final.bold = True

doc.add_paragraph()
doc.add_paragraph()

# 푸터
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("─" * 30)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f2 = footer2.add_run("READING JESUS DAILY Q.T.")
run_f2.italic = True

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f3 = footer3.add_run("매일, 삶 속에서 예수님을 읽어갑니다.")
run_f3.italic = True

footer4 = doc.add_paragraph()
footer4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f4 = footer4.add_run("말씀으로 365 — 읽는 교회를 넘어, 말씀과 함께 사는 교회로.")
run_f4.italic = True

# 저장
doc.save('C:/Lacal_workspace/project/reading-jesus/docs/QT/READING_JESUS_QT_가이드.docx')
print("Word 문서가 생성되었습니다!")
